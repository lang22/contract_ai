#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Time       : 2019/2/26 9:38
@Author     : jzs
@File       : __init__.py.py
@Software   : PyCharm
@Description: 新的合同条款审核

合同条款审核条款句子以及条款集合的定义
非条款的界限：开头是模式为“第一条 **** ”句子之前的句子，可以认为是一个特殊的大条款
大条款的界限：开头是模式为“第N条 **** ”句子, 结束为但不包含模式为“第N+1条 ****”的句子的若干个子条款集合
子条款的界限：开头是模式“n.m **** ”句子, 结束为但不包含模式为“n.m+1 ****”的句子的若干个句子（条款句）集合
条款句子：最基础的句子

由定义出发，我们需要定义界限特征非常明显的条款句子以及条款集合，即，大条款和条款句子的类。
在这两个基础类建立各种类的操作以及属性


"""
from docx import Document
from app.tools import random_fliename
from config import Config
from app.tools import join_path

from app.tools.xlsx_tools import gen_xlsx
from .exam_content import get_main_clauses_paragraphs, set_no_clauses_sentence_matching, \
    set_clauses_list_matching_clauses, set_exam_diff_result, exam_result_to_html, exam_result_to_xlsx, \
    exam_result_to_docx
from .cont_sentences import Sentence, DEFAULT_SENTENCE_SIMILARITY_MODEL, \
    NO_CLAUSE_SENTENCE_SIMILARITY_MODEL


def get_exam_result_by_html(template_path: str, test_path: str) -> str:
    """
    输入两个文档，得到两个文档的审核结果，
    将两个文档分别切分成条款和非条款，
    分别通过语义设置两个文档的条款和非条款的最匹配句，
    审核结果用带有样式的html代码(str)表示。

    :param template_path: 模板文档的路径
    :param test_path:  待审核文档的路径
    :return:
    """
    # 将两个文档分别切分成条款和非条款
    template_document = Document(template_path)
    test_document = Document(test_path)

    template_no_clause, template_clause_list = get_main_clauses_paragraphs(document=template_document, document_id=1)
    test_no_clause, test_clause_list = get_main_clauses_paragraphs(document=test_document, document_id=2)

    # 分别通过语义设置两个文档的条款和非条款的最匹配句
    set_no_clauses_sentence_matching(template_no_clause, test_no_clause)
    set_clauses_list_matching_clauses(template_clause_list, test_clause_list)

    # 根据匹配句子结果进行差异对比
    set_exam_diff_result(first_no_clauses=template_no_clause,
                         second_no_clauses=test_no_clause,
                         first_clauses_list=template_clause_list,
                         second_clauses_list=test_clause_list)

    # 将审核结果转化成html字符串
    template_html = exam_result_to_html(template_no_clause, template_clause_list)
    test_html = exam_result_to_html(test_no_clause, test_clause_list)

    # 保存审核结果到审核结果返回缓存队列中
    other_result_dict = {
        'template_document': template_document,  # 模板文档对象
        'template_no_clause': template_no_clause,  # 模板非条款对象
        'template_clause_list': template_clause_list,  # 模板条款对象集合

        'test_document': test_document,  # 测试文档对象
        'test_no_clause': test_no_clause,  # 测试非条款对象
        'test_clause_list': test_clause_list  # 测试条款对象集合
    }

    # 生成docx文档1
    file_path1 = join_path(Config.OTHER_DOWNLOAD_PATH, "合同文档1-" + random_fliename() + ".docx")
    download_result_docx(other_result_dict, file_path1, True)

    # 生成docx文档2
    file_path2 = join_path(Config.OTHER_DOWNLOAD_PATH, "合同文档2-" + random_fliename() + ".docx")
    download_result_docx(other_result_dict, file_path2, False)

    # 生成xlsx文档
    xlsx_download_path = join_path(Config.OTHER_DOWNLOAD_PATH, "合同文档对比结果" + random_fliename() + ".xlsx")
    download_result_xlsx(other_result_dict, xlsx_download_path)

    return template_html, test_html, file_path1, file_path2, xlsx_download_path


def get_exam_result_by_html2(template_path: str, test_path: str) -> str:
    """
    输入两个文档，得到两个文档的审核结果，
    将两个文档分别切分成条款和非条款，
    分别通过语义设置两个文档的条款和非条款的最匹配句，
    审核结果用带有样式的html代码(str)表示。

    :param template_path: 模板文档的路径
    :param test_path:  待审核文档的路径
    :return:
    """
    # 将两个文档分别切分成条款和非条款
    template_document = Document(template_path)
    test_document = Document(test_path)

    template_no_clause, template_clause_list = get_main_clauses_paragraphs(document=template_document, document_id=1)
    test_no_clause, test_clause_list = get_main_clauses_paragraphs(document=test_document, document_id=2)

    # 分别通过语义设置两个文档的条款和非条款的最匹配句
    set_no_clauses_sentence_matching(template_no_clause, test_no_clause)
    set_clauses_list_matching_clauses(template_clause_list, test_clause_list)

    # 根据匹配句子结果进行差异对比
    set_exam_diff_result(first_no_clauses=template_no_clause,
                         second_no_clauses=test_no_clause,
                         first_clauses_list=template_clause_list,
                         second_clauses_list=test_clause_list)

    # 将审核结果转化成html字符串
    template_html = exam_result_to_html(template_no_clause, template_clause_list)
    test_html = exam_result_to_html(test_no_clause, test_clause_list)

    # 保存审核结果到审核结果返回缓存队列中
    other_result_dict = {
        'template_document': template_document,  # 模板文档对象
        'template_no_clause': template_no_clause,  # 模板非条款对象
        'template_clause_list': template_clause_list,  # 模板条款对象集合

        'test_document': test_document,  # 测试文档对象
        'test_no_clause': test_no_clause,  # 测试非条款对象
        'test_clause_list': test_clause_list  # 测试条款对象集合
    }

    # 生成docx文档1
    file_path1 = join_path(Config.OTHER_DOWNLOAD_PATH, "合同文档1-" + random_fliename() + ".docx")
    download_result_docx(other_result_dict, file_path1, True)

    # 生成docx文档2
    file_path2 = join_path(Config.OTHER_DOWNLOAD_PATH, "合同文档2-" + random_fliename() + ".docx")
    download_result_docx(other_result_dict, file_path2, False)


    return template_html, test_html, file_path1, file_path2


def download_result_xlsx(other_result_dict, download_path: str):
    """
    将审核结果标注成xlsx文件

    :param other_result_dict: : 审核结果
    :param download_path:下载路径
    :return:
    """
    clause_list1: list = other_result_dict['template_clause_list']
    clause_list2: list = other_result_dict['test_clause_list']
    no_clause1: list = other_result_dict['template_no_clause']
    no_clause2: list = other_result_dict['test_no_clause']

    new_result_list1, new_result_list2 = exam_result_to_xlsx(clause_list1, clause_list2, no_clause1, no_clause2)
    gen_xlsx(new_result_list1, new_result_list2, download_path)


def download_result_docx(other_result_dict, download_path: str, is_first: bool) -> bool:
    """
    标注docx文件并保存

    :param other_result_dict: 审核结果
    :param download_path: 下载路径
    :param is_first: 是否是模板文档
    :return:
    """

    if is_first:
        download_document = other_result_dict['template_document']
        no_clauses = other_result_dict['template_no_clause']
        clauses_list = other_result_dict['template_clause_list']
    else:
        download_document = other_result_dict['test_document']
        no_clauses = other_result_dict['test_no_clause']
        clauses_list = other_result_dict['test_clause_list']

    return exam_result_to_docx(download_document,
                               no_clauses,
                               clauses_list,
                               download_path)
