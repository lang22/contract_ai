#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Time       : 2019-04-09 10:22
@Author     : charles
@File       : solve_html.py
@Software   : PyCharm
@Description: ......
"""
import re
from pydocx import PyDocX

from docx import Document

from cont_check.ES_check_api import get_sims_by_ES


def get_content_from_string(string: str):
    """
    从p标签中获取内容
    :param string:
    :return:
    """
    start_index = -1
    end_index = -1
    for i, ch in enumerate(string):
        if i == len(string) - 1:
            break
        if ch == '>' and string[i + 1] != '<':
            start_index = i + 1
        if ch != '>' and string[i + 1] == '<':
            end_index = i

    return string[start_index:end_index + 1]


def solve_html(html):
    """
    修改插件生成的html，生成正确格式
    :param html_name:
    :return:
    """
    res = re.findall('<p>.*?</p>', html, re.S)
    for item in res:

        if item.find('class="pydocx-right"') >= 0:
            new_item = item.replace('<p>', '<p class="pydocx-right">')
            html = html.replace(item, new_item)

        if item.find('class="pydocx-center"') >= 0:
            new_item = item.replace('<p>', '<p class="pydocx-center">')
            html = html.replace(item, new_item)

    return html


def to_html(filepath):
    """
    根据文件路径将.docx生成html文件，命名为test.html
    :param filepath:
    :return:
    """
    html = PyDocX.to_html(filepath)
    return html


def get_body_css(html):
    """
    将html内容中的body与style返回
    :param filepath:
    :return:
    """
    style_res = re.findall('<style>.*?</style>', html, re.S)[0]
    body_res = re.findall('<body>.*?</body>', html, re.S)[0]
    style_res = re.sub('body.*?}', '', style_res)
    return body_res, style_res


def add_link(html):
    """
    把html中的p标签中的内容，如果存在相似信息，加上a标签
    :param html_name:
    :return:
    """
    res = re.findall('<p>.*?</p>|<li>.*?</li>', html, re.S)
    sims_lis = []
    for i, item in enumerate(res):
        para_content = get_content_from_string(item)

        doc = Document()
        doc.add_paragraph(para_content)
        save_name = 'temp.docx'
        doc.save(save_name)

        response_text = get_sims_by_ES(save_name) # todo 此处为运行相似度的位置


        if 'sims' in response_text:  # 存在相似信息
            sims_lis.append(response_text['sims'])
            new_item = '<a class="sims_item" onclick="javascript:test(\'' + para_content + '\')" >' + item + '</a>'
            html = html.replace(item, new_item)
        else:  # 不存在相似信息
            sims_lis.append('无相似信息')

    return html, sims_lis


def solve_method(filepath):
    """
    将文档返回html的body与style，以及第一个可点击内容的相似信息
    :param filepath:
    :return:
    """
    html = to_html(filepath)
    html = solve_html(html)
    html, sims_lis = add_link(html)  # sims_lis: 每一段的相似信息 eg: ['无相似信息', [{}, {}] ]

    # 找出第一个有相似信息的段落
    for i, item in enumerate(sims_lis):
        if item != '无相似信息':
            first_sims = item  # 记录第一个有相似信息的段落的相似信息
            break

    if (i + 1) == len(sims_lis):
        first_sims = []

    # 获取这个段落的段落内容
    doc_temp = Document(filepath)
    para_list = list(doc_temp.paragraphs)
    para_content = para_list[i].text

    for item in first_sims:

        simple_name = item['sim_documents_path']
        if len(simple_name) >= 20:
            item['simple_name'] = item['sim_documents_path'][:20] + '...'
        else:
            item['simple_name'] = item['sim_documents_path']
        item['sims'] = str(item['sims'])[:-11]

        # print(type(item))

    body, style = get_body_css(html)
    return body, style, first_sims, para_content





if __name__ == '__main__':
    filepath = "/Users/charles/bupt-file/合同机器人/条款审核测试文档/123.docx"
    body, style, first_sims = solve_method(filepath)
