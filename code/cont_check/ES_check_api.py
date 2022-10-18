from collections import Counter

import requests
from docx import Document
import json

from app.models.flower_table import Flower
from app.tools.aes import Prpcrypt
from cont_check.singleF import sim_interface
from config import Config

ES_interface1_url = "http://10.64.140.116:8080/fullTextSearch"  # 获取合规内容接口
ES_interface2_url = "http://10.64.140.116:8080/SingleF"  # 获取相似信息接口


def sort_sims_file(results):
    """
    将生成的results中的sims信息重新排序
    :param results: [ {'paragraph':'...', 'yellow':'...', 'sims':[{}. {} ....] } ]
    :return:
    """
    result = results[0]
    if 'sims' in result:
        sims = result['sims']

        name_info_dic = {dic['sim_documents_path']: dic for dic in sims}  # {name: dict字典信息}

        cont_name_lis = [dic['sim_documents_path'] for dic in sims]  # 记录下所有的相似合同名称

        paragraph = result['paragraph']
        aes = Prpcrypt(Config.DB_STR_AES_KEY)
        aes_cont_content = aes.encrypt(paragraph)

        # 将对点赞过的相似信息合同名称记录下来
        cont_clues = []

        elements = Flower.get_item_by_para_content(aes_cont_content)
        for element in elements:
            cur_name = element.sim_cont_name
            decrypt_name = aes.decrypt(cur_name).replace('\x00', '')
            cont_clues.append(decrypt_name)
        # 优先排序
        c = Counter(cont_clues)
        commons = c.most_common()

        sort_results = []

        for com in commons:
            cur_cont_name = com[0]
            sort_results.append(name_info_dic[cur_cont_name])
            cont_name_lis.remove(cur_cont_name)

        for item in cont_name_lis:
            sort_results.append(name_info_dic[item])

        result['sims'] = sort_results

        return result
    else:
        return result


# def get_fullText_by_ES(filename):
#     """
#     通过文件名调用ES接口获取文件内容  公司系统ES接口
#     :param filename: 文件名
#     :return:
#     """
#     querystring = {"title": filename}
#
#     payload = ""
#     headers = {
#         'Content-Type': "application/json",
#         'cache-control': "no-cache",
#         'Postman-Token': "39ac530f-b4c7-4c13-b511-cea02d7f7685"
#     }
#
#     response = requests.request("GET", ES_interface1_url, data=payload, headers=headers, params=querystring)
#
#     all_content = json.loads(response.text)
#     # 内容包含 title fullText
#     title = all_content['title']
#     fullText = all_content['fullText']
#     print('fullText: ', fullText)
#     return title, fullText


with open(Config.LAW_LIBRARY_DATA_PATH, 'r', encoding='utf-8') as f:
    j_dict = json.load(f)


def get_fullText_by_ES(filename):
    """
    通过文件名调用ES接口获取文件内容  本地json文件
    :param filename: 文件名
    :return:
    """
    global j_dict

    fullText = j_dict[filename]['_source']['additionFileds']['fullText']
    title = j_dict[filename]['_source']['title']
    return title, fullText.split('<br>')


# def get_sims_by_ES(filename):
#     """
#     通过文件名调用ES接口，获取合同每一段的相似信息  公司系统接口模式
#     :param filename: 文件名
#     :return:
#     """
#     files = {'file': open(filename, 'rb')}
#     response = requests.post(ES_interface2_url, files=files)
#     return response.text


def get_sims_by_ES(filename):
    """
    通过文件名调用ES接口，获取合同每一段的相似信息   本地数据模式
    :param filename: 文件名
    :return:
    """
    paragraphs = []
    f = Document(filename)
    for i in range(len(f.paragraphs)):
        text = f.paragraphs[i].text
        if text:
            paragraphs.append(text)

    if paragraphs:
        results = sim_interface(paragraphs)
        result = sort_sims_file(results)

        return result
    else:
        return []


if __name__ == '__main__':
    file = '/Users/charles/Desktop/old合同.docx'
    file2 = '/Users/charles/bupt-file/合同机器人/条款审核测试文档/债权转让协议_打包收购 - 条款0.docx'

    doc_temp = Document(file)
    paragraphs = doc_temp.paragraphs

    for item in paragraphs:
        print(item.text)
        print('--------------')
