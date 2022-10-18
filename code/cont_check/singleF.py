import os

import docx
import timeit
import jieba
# jieba.load_userdict("data.txt") #这是一个自定义词典库
from gensim import corpora, models, similarities
from gensim import corpora
import pickle
import pandas as pd
from config import Config

dictionary = corpora.Dictionary.load(os.path.join(Config.LSI_MODEL_PATH, "paper_dict.dict"))
index = similarities.MatrixSimilarity.load(os.path.join(Config.LSI_MODEL_PATH, 'lsi_model.index'))
pkl_file = open(os.path.join(Config.LSI_MODEL_PATH, 'documents_key.pkl'), 'rb')
documents_key = pickle.load(pkl_file)
pkl_file.close()
# corpus = corpora.MmCorpus('paper_corp.mm')
lsi = models.LsiModel.load(os.path.join(Config.LSI_MODEL_PATH, 'model.lsi'))
pkl_file2 = open(os.path.join(Config.LSI_MODEL_PATH, 'documents_whole.pkl'), 'rb')
documents_whole = pickle.load(pkl_file2)
pkl_file2.close()
start = timeit.default_timer()


def split(text):
    """直接将text文本进行切分，返回由词构成的list"""
    words = jieba.cut(text, HMM=True, cut_all=False)
    return words


def stopwordslist(filepath):
    """创建停用词表"""
    stopwords = [line.strip() for line in open(filepath, 'r', encoding='gbk').readlines()]
    return stopwords


# 去除停止词和非中文
stop_words = stopwordslist(os.path.join(Config.LSI_MODEL_PATH, 'stop_words.txt'))


def is_uchar(uchar):
    """判断一个unicode是否是汉字"""
    if uchar >= u'\u4e00' and uchar <= u'\u9fa5':
        return True
    return False


def read_data(path):
    """从文档中读取数据"""
    paragraphs = []
    f = docx.Document(path)
    for i in range(len(f.paragraphs)):
        paragraphs.append(f.paragraphs[i].text)
    return paragraphs


def cal_Sim(text):
    split_text = list(split(text))
    new_text = []
    for i in split_text:
        if not len(i) == 0 and is_uchar(i):
            if i not in stop_words:
                new_text.append(i)

    if len(new_text) < 10:
        # print('段落过短，无法判断相似信息')
        return 0
    new_vec = dictionary.doc2bow(new_text)

    vec_lsi = lsi[new_vec]

    sims = index[vec_lsi]
    sims = sorted(enumerate(sims), key=lambda item: -item[1])

    if sims[0][1] < 0.8:
        # print('相似度过低')
        return 0

    result = []
    for i in sims[0:10]:
        result.append([documents_key[i[0]][0], documents_whole[i[0]], i[1]])

    return result


def sim_interface(paragraphs: "list[str]"):
    df_yellow = pd.DataFrame(columns=["paragraph", "yellow"])
    df_sims = pd.DataFrame(columns=["paragraph_num", "sim_documents_path", "sim_document", "sims"])
    new_df_yellow = []
    for i in range(len(paragraphs)):
        result = cal_Sim(paragraphs[i])

        if result:
            new_df_sims = []
            for single_result in result:
                address = single_result[0].split('files\\')
                address = address[1][:-5]
                new_df_sims.append({"paragraph_num": i, "sim_documents_path": address, "sim_document": single_result[1],
                                    "sims": float(single_result[2])})

            new_df_yellow.append({"paragraph": paragraphs[i], "yellow": 1, "sims": new_df_sims})
        else:
            new_df_yellow.append({"paragraph": paragraphs[i], "yellow": 0})
            df_yellow = df_yellow.append(new_df_yellow, ignore_index=True)
    return new_df_yellow


if __name__ == '__main__':
    text = ['保密期限：在本合同有效期内以及在有效期后的合理期限内甲乙双方均不得将合同及项目相关的技术资料、技术秘密等成为公共信息之前披露给任何第三方。']
    result = sim_interface(text)

    for item in result:
        print(item)
