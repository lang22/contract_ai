import copy
import time
from docx import Document

from app.json_models.admin_check_search_json_result import parse_json
from app.models.cont_check_table import ContCheckTable
from app.tools import join_path, random_fliename
from app.tools.aes import Prpcrypt
from config import Config
from cont_check.ES_check_api import get_sims_by_ES, sort_sims_file
from cont_check.singleF import sim_interface


def get_check_results(filename: str, para_index: int) -> list:
    """
    获取合同文档中某一段的相似信息结果
    :param filename: 文件名
    :param para_index: 段落号
    :return: 结果词典
    """
    para_temp = Document(filename)
    results = list()
    for item in para_temp.paragraphs:
        text = item.text
        if text:
            results.append(text)

    sentence = results[para_index]
    doc = Document()
    doc.add_paragraph(sentence)
    save_name = join_path(Config.UPLOADED_DIR_PATH, random_fliename() + '.docx')
    doc.save(save_name)

    response_text = get_sims_by_ES(save_name)

    result = parse_json(response_text)

    return result


def save_to_mysql(cont_name: str, content_list: list):
    """
    将上传文档上传至数据库 filename content user_id upload_time
    文档内容以及文档名需要加密之后进行存储
    :param filename: 1.docx 格式
    :param content_list:
    :return:
    """
    try:
        upload_time = time.strftime('%Y.%m.%d', time.localtime(time.time()))

        content = ''.join(content_list)

        aes = Prpcrypt(Config.DB_STR_AES_KEY)

        aes_content = aes.encrypt(content)
        aes_cont_name = aes.encrypt(cont_name)

        ContCheckTable.add_one(cont_name=aes_cont_name, upload_time=upload_time, cont_content=aes_content)
        print('合规性审核合同存储数据库成功')

    except Exception as e:
        print(e)
        print('合规性审核合同存储数据库失败')


def get_sims_result_by_sentence(sentence: str):
    """
    根据句子获得相似信息
    :param sentence:
    :return:
    """
    doc = Document()
    doc.add_paragraph(sentence)
    save_name = 'temp.docx'
    doc.save(save_name)

    response_text = get_sims_by_ES(save_name)
    return response_text


def sims_result_parse_json(text: str):
    """
    将数据处理成我们需要的json格式  本地数据模式
    :param text:
    :return:
    """
    final_result = {'success': '0', 'results': ''}
    results = list()
    dic = text

    if 'sims' in dic:
        # 这一段存在相似段落信息
        final_result['success'] = '1'
        sims_dic = dic['sims']

        for item in sims_dic:
            filename = item['sim_documents_path']
            if len(filename) >= 20:
                simple_file = filename[:20] + '...'
            else:
                simple_file = filename

            cur_temp = {'filename': filename, 'content': item['sim_document'],
                        'sims': str(item['sims'])[:-11], 'simple_file': simple_file}
            results.append(cur_temp)
    else:
        # print('本段没有相似段落信息')
        pass

    final_result['results'] = results

    return final_result


def get_sims_result_by_sentence(content: str) -> list:
    """
    通过一个句子获取相似度结果
    :param content:
    :return:
    """
    results = sim_interface([content])
    result = sort_sims_file(results)
    return result


def get_sims_result_by_paragraphs(paragraphs: 'list[str]') -> list:
    """
    获取多个字符串段落的相似信息
    :param paragraphs:
    :return:
    """
    if not paragraphs:
        return []
    sims_list = []
    item_tmp = {"paraID": '', "paraContent": '', "clickable": 1}

    find_first = False
    first_content_result = None
    for i, para in enumerate(paragraphs):
        item = copy.deepcopy(item_tmp)
        # 寻找第一个有相似信息的
        result = get_sims_result_by_sentence(para)
        if not find_first and 'sims' in result:
            find_first = True
            first_content_result = result

        item['clickable'] = 1 if 'sims' in result else 0
        item['paraID'] = 'paraID' + str(i)
        item['paraContent'] = para
        sims_list.append(item)

    return first_content_result, sims_list
