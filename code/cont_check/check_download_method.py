from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from config import Config
from cont_check.ES_check_api import get_fullText_by_ES

url = "http://10.64.140.116:8080/fullTextSearch"


def get_content_by_filename(filename: str, match_content: str):
    """
    通过文件名调用ES接口获得整篇文档内容，相似内容在文档中高亮显示为蓝色
    :param filename:
    :return:
    """

    title, fullText = get_fullText_by_ES(filename)

    lis = list()
    lis.append(title)

    lis.extend(item for item in fullText)

    fontsize = 14.0
    fontname = u'宋体'

    doc = Document()
    # doc.add_heading(title, 0)

    for sentence in lis:
        paragraph = doc.add_paragraph('')

    for index, paragraph in enumerate(doc.paragraphs):
        run = paragraph.add_run(lis[index])
        run.font.size = Pt(fontsize)
        run.font.name = fontname
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        if lis[index] == match_content:
            run.font.highlight_color = getattr(WD_COLOR_INDEX, "YELLOW")

    doc.save(Config.DOWNLOAD_PATH + '/' + title)






