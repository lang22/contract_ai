#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Author     : cml
@File       : __init__.py
@Software   : PyCharm
@Description: 合同生成主调模块
"""

import os
import re

from docx import Document
from docx.oxml.ns import qn

from app.tools import random_fliename
from config import Config
from cont_gen.base_tools import trans_date
from cont_gen.solve_method import solve_single_instruction, solve_date, solve_case_ammount, solve_table4

template_path = Config.CONT_GENA_TEMPLATE_DIR_PATH

# 新生成的模板文档保存路径
filled_docx_path = Config.OTHER_DOWNLOAD_PATH

# 处理不同生成方式的函数字典
SOLVE_METHOD = {
    'simple_instructions': solve_single_instruction,
    'date': solve_date,
    'case_amount': solve_case_ammount,
    'table4': solve_table4
}


def __read(filepath: str):
    """
    读取整篇文章，分段存储

    :param filepath: 文件路径
    :return:
    """
    doc_temp = Document(filepath)

    # para_list = list(doc_temp.paragraphs)

    return doc_temp


def contract_generate_func(choose_template_name: str,
                           contract_filename: str,
                           elem_content_dict: dict,
                           elem_id_dict: dict,
                           elem_key_dict: dict,
                           elem_gena_type_dict: dict):
    """
    合同生成主调函数

    输入用户选的模版名称、用户另存为的文件名和用户输入的value值（要素-用户输入的value值字典）
    根据要素-要素ID字典、要素-关键字字典和要素-生成类型字典查找模板文件中填写要素的相应位置，并基于规则填写

    :param choose_template_name: 用户选的模版名称
    :param elem_content_dict: 用户输入的value值（要素-用户输入的value值字典）
    :param contract_filename: 用户另存为的文件名
    :param elem_id_dict: 要素-要素ID字典
    :param elem_key_dict: 要素-关键字字典
    :param elem_gena_type_dict: 要素-生成类型字典
    :return:
    """
    template_docx_filename = choose_template_name + '.docx'
    template_docx_path = os.path.join(template_path, template_docx_filename)  # 要拿到的转让协议的路径

    doc_temp = __read(template_docx_path)  # 读取模版文档

    # 段落
    para_list = list(doc_temp.paragraphs)

    for item in elem_id_dict:

        keys = elem_key_dict[item]  # elem_key 字段
        tar = elem_content_dict[item]  # 要填充的目标值
        instruction = elem_gena_type_dict[item]  # 操作指令类型

        if keys == ['', ''] or keys == '':
            continue

        if tar == '' or tar == ',':  # 如果用户没有对这个字段进行填写，不对这个字段进行处理
            continue

        if instruction == 'case_amount':
            tar = tar[0] + ',' + tar[1]

        func = SOLVE_METHOD[instruction]

        try:
            if instruction == 'table4':
                row_index = keys[4]
                func(doc_temp, tar, int(row_index))
            else:
                func(para_list, keys, tar, instruction)

        except Exception as e:
            print('数据库中无此选项', item, 'error:', e)
            continue

    filled_docx_filename = contract_filename + random_fliename() + '.docx'
    save_filled_docx_path = os.path.join(filled_docx_path, filled_docx_filename)
    doc_temp.styles['Normal'].font.name = u'宋体'
    doc_temp.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc_temp.save(save_filled_docx_path)

    return filled_docx_filename
