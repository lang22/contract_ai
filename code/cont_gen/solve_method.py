import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from cont_gen.para_index_compute import get_para_index
from cont_gen.base_tools import REGEX_EXPRESSION, trans_date


# def __add_text(paragraph, text):
#     """
#     在段落中添加文字
#
#     :param paragraph: 段落
#     :param text: 文字
#     :return:
#     """
#     font = paragraph.runs[0].font
#     paragraph.text = ''
#     run = paragraph.add_run(text)
#     run.font.size = font.size
#     run.font.name = font.name
#
#     return run


def __add_text(paragraph, text):
    """
    在段落中添加文字

    :param paragraph: 段落
    :param text: 文字
    :return:
    """
    fontsize = 14.0
    fontname = u'宋体'
    for i in range(len(paragraph.runs)):
        paragraph.runs[i].clear()
    run1 = paragraph.add_run(text)
    run1.font.size = Pt(fontsize)
    run1.font.name = fontname


def solve_single_instruction(para_list: list,
                             keywords: str,
                             tar_key: str,
                             type: str):  # single_instruction 类型的处理

    """

    :param doc_temp: 需要操作的文件
    :param keywords 关键词串，注意考虑到多字段以及同义词
    :param tar_key  目标替换串
    :param type: 操作指令类型
    :return:
    """

    para_index, final_key = get_para_index(para_list, keywords, type)
    tmp = para_list[para_index].text.replace(' ', '')

    final_key_index = tmp.find(final_key)
    target_index = tmp[final_key_index:].find('[]')
    text = tmp[:final_key_index] + tmp[final_key_index:].replace('[]', '[%s]' % tar_key, 1)

    if keywords == '编号':
        text = '                                  ' + text

    __add_text(para_list[para_index], text)


def solve_date(para_list, keywords: str, tar_key: str, type: str):  # date 类型的处理

    """

    :param doc_temp: 需要操作的文件
    :param keywords 关键词串，注意考虑到多字段以及同义词
    :param tar_key  目标替换串
    :param type: 操作指令类型
    :return:
    """

    para_index, key = get_para_index(para_list, keywords, type)
    date_list = trans_date(tar_key)

    tmp = para_list[para_index].text.replace(' ', '')

    text = tmp.replace('[]年[]月[]日', '[%s]年[%s]月[%s]日' % (date_list[0], date_list[1], date_list[2]), 1)

    __add_text(para_list[para_index], text)


def solve_case_ammount(para_list, keywords: str, tar_key: str, type: str):
    """

    :param doc_temp: 需要操作的文件
    :param keywords 关键词串，注意考虑到多字段以及同义词
    :param tar_key  目标替换串
    :param type: 操作指令类型
    :return:
    """
    para_index, key = get_para_index(para_list, keywords, type)

    UPPER = tar_key.split(',')[0]
    LOWWER = tar_key.split(',')[1]
    restr = REGEX_EXPRESSION[type] % key

    tmp = para_list[para_index].text.replace(' ', '')
    res = re.match(restr, tmp).group(1)
    index = tmp.find(res)

    first_section = tmp[0:index]  # 对文件进行修改
    second_section = tmp[index:].replace('[]', '[%s]', 2) % (UPPER, LOWWER)
    text = ''.join([first_section, second_section])

    __add_text(para_list[para_index], text)


def solve_table4(document, target: str, row_index: int):
    """
    填写表格内容，table4格式
    :param document: 读取文件后的document
    :param target: 要填写的内容
    :param index: 要填写表格的row值
    :return:
    """
    try:
        tables = document.tables

        if len(tables) < 2:
            return

        table = tables[0]
        tar_list = target.split('，')

        for i in range(3):
            tar_key = tar_list[i]
            run = table.cell(row_index, i+1).paragraphs[0].add_run(tar_key)
            run.font.name = u'宋体'
            run.font.size = Pt(14.0)
            table.cell(row_index, i+1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    except Exception as e:
        print(e)
        print('输入的内容不符合表格填写格式')
