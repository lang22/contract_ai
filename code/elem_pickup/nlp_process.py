#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Time       : 2018/9/27 9:43
@Author     : jzs
@File       : nlp_main.py
@Software   : PyCharm
@Description: 主要用于文档的一些预处理和语料库的建立
"""
from collections import defaultdict
from functools import reduce
from typing import Dict
from typing import List
from typing import Tuple
from gensim.models import word2vec, Word2Vec
import pickle
import docx
import jieba
import os
import re

# 停用词列表文件路径
base_path = os.path.dirname(__file__)
STOP_WORD = os.path.join(base_path, 'stop_word.model')
WORD2VEC_MODEL = os.path.join(base_path, 'word2vec.model')

# 默认标点词列表文件路径
DEFAULT_PUNCTUATION = '。？！，、；：“”‘’（）─……—·《》〈〉/_【】'

# 正则表达式特殊字符转义列表
RE_CHAR_LIST = [
    ("\\", "\\\\"), ("*", "\\*"),
    ("+", "\\+"), ("|", "\\|"),
    ("{", "\\{"), ("}", "\\}"),
    ("(", "\\("), (")", "\\)"),
    ("^", "\\^"), ("$", "\\$"),
    ("[", "\\["), ("]", "\\]"),
    ("?", "\\?"), (",", "\\,"),
    (".", "\\."), ("&", "\\&")
]


def re_char_dict() -> 'Dict[str,int]':
    """
    得到正则表达式特殊字符转义字典
    :return:
    """
    re_dict = dict()
    for i in range(len(RE_CHAR_LIST)):
        key = RE_CHAR_LIST[i][0]
        re_dict[key] = i
    return re_dict


class NLPProcess(object):
    """
    主要用于文档的一些预处理和语料库的建立
    """
    # 停用词字典
    STOP_WORD_DIC = None

    # 标点查找字符串
    PUNCTUATION_STR = None

    # 语料库（词库），如['hello','word',...,'python']
    @property
    def corpus(self) -> 'List[str]':
        return self.__corpus

    # 所有句子集合, 如['hello word !', 'I love python!']
    @property
    def sentences(self) -> 'List[str]':
        return self.__sentences

    # 分词后的句子的句子集合，如[['hello','word'], ['I','love', 'python']]
    @property
    def words_sentences(self):
        return self.__words_sentences

    # 特征向量索引字典，如eigenvector_dict['word']=0.0
    @property
    def eigenvector_dict(self) -> 'Dict[str,float]':
        return self.__eigenvector_dict

    # word2vec词向量索引字典，如word2vec_dict['word']=[0.0,0.0,...,0.0]
    @property
    def word2vec_dict(self) -> 'Dict[str,List[float]]':
        return self.__word2vec_dict

    def __init__(self, path: str, isfile: bool, *punctuation_list: 'Tuple[str]'):
        """
        构造器，将导入文档，文档预处理，建立语料库，并基于语料库中的词建立特征向量索引字典

        :param path: 文件径
        :param isfile: 是否路是单个文件
        :param punctuation_list: 选择切分的符号
        """
        # 导入停用词字典、标点查找字符串
        NLPProcess.STOP_WORD_DIC = NLPProcess.__load_stop_word_file_to_dic(STOP_WORD)
        NLPProcess.PUNCTUATION_STR = NLPProcess.__load_punctuation(*punctuation_list) \
            if len(punctuation_list) else DEFAULT_PUNCTUATION

        tmp = NLPProcess.__get_paragraphs(path, isfile)

        self.__sentences = NLPProcess.__paragraphs_to_sentences(tmp)
        self.__corpus = self.__create_corpus()
        self.__eigenvector_dict = self.__create_eigenvector_dict()
        self.__words_sentences = self.__get_words()

        self.__word2vec_dict = NLPProcess.__words_to_vector(self.__words_sentences)

    def __create_eigenvector_dict(self) -> Dict[str, int]:
        """
        基于语料库中的词建立特征向量索引字典
        :return:
        """
        index_list = [i for i in range(len(self.corpus))]
        return dict(zip(self.corpus, index_list))

    def __create_corpus(self) -> 'Set[str]':
        """
        建立语料库，将所有段落按照规定标点符号，分割成句子；
        通过进行预处理分词，加入集合去除重复的词；
        排序确定语料库中的词的顺序
        :return:
        """
        words_list = NLPProcess.sentences_to_words(self.__sentences)
        tmp_list = list()

        for words in words_list:
            tmp_list.extend(words)

        tmp_list = list(set(tmp_list))

        return sorted(tmp_list)

    def __get_words(self) -> 'List[List[str]]':
        """
        将段落按照规定标点符号，分割成句子，然后通过进行预处理分词
        :return:
        """
        tmp = NLPProcess.__paragraphs_to_sentences(self.sentences)
        return NLPProcess.sentences_to_words(tmp)

    @staticmethod
    def __get_paragraphs(path: str, isfile: bool) -> 'List[str]':
        """
        从文档或文件夹集合读取读取段落
        :param path:文件路径
        :param isfile:
        :return:
        """
        sentences = NLPProcess.__load_docx_file_paragraphs(path) \
            if isfile else NLPProcess.__load_docx_dir_paragraphs(path)
        return sentences

    @staticmethod
    def get_paragraphs(path: str, isfile: bool) -> 'List[str]':
        """
        从文档或文件夹集合读取读取段落
        :param path:文件路径
        :param isfile:
        :return:
        """
        sentences = NLPProcess.__load_docx_file_paragraphs(path) \
            if isfile else NLPProcess.__load_docx_dir_paragraphs(path)
        return sentences

    @staticmethod
    def __load_docx_file_paragraphs(file_path: str) -> 'List[str]':
        """
        读取一个docx文档每一段落,并得到该字符串对象,使用列表存储
        :param file_path: 文件路径
        :return: 文档每一段落字符串对象列表
        """
        try:
            docx_file = docx.Document(file_path)
            return [para.text for para in docx_file.paragraphs if para.text != '']
        except docx.opc.exceptions.PackageNotFoundError as e:
            print(e)
            return list()

    @staticmethod
    def __load_docx_dir_paragraphs(dir_path: str) -> 'List[str]':
        """
        读取一个文件夹下的所以docx文档每一段落,合并成列表，并得到该字符串对象,使用列表存储
        :param dir_path: 文件路径
        :return: 文档每一段落字符串对象列表
        """
        try:
            file_path_set = os.listdir(dir_path)
            result = list()
            for file_path in file_path_set:
                docx_file = docx.Document(dir_path + '/' + file_path)
                result.extend([para.text for para in docx_file.paragraphs if para.text != ''])
            return result
        except BaseException as e:
            print(e)
            return list()

    @staticmethod
    def __load_stop_word_file_to_dic(file_path: str) -> 'Dict[str, int]':
        """
        读取停用次列表文件，并使用字典存储
        :param file_path: 文件路径
        :return: 停用次字典
        """
        with open(file_path, "r", encoding='utf-8') as file:
            list1 = [line.strip() for line in file.readlines()]
            list2 = [i for i in range(len(list1))]
            stop_word_dic = dict(zip(list1, list2))
            return stop_word_dic

    @staticmethod
    def __load_punctuation_file_to_re_str(file_path: str) -> str:
        """
        读取标点符号文件，并保存为正则表达式字符串
        :return:
        """
        with open(file_path, "r", encoding='utf-8') as file:
            string = file.readline()
            re_v = reduce(lambda x, y: x + '|' + y, [v[1] for v in RE_CHAR_LIST])
            punctuation_str = reduce(lambda x, y: x + '|' + y, [s for s in string])
            punctuation_str = punctuation_str + '|' + re_v
            return punctuation_str

    @staticmethod
    def __load_punctuation(*punc_list: List[str]) -> str:
        """
        将参数中的标点符号列表，若有正则符号，先将其转换，然后再转换成正则字符串
        :return:
        """
        tl = [x for x in punc_list]
        re_dict = re_char_dict()
        for i in range(len(tl)):
            value = re_dict.get(tl[i])
            if value:
                tl[i] = RE_CHAR_LIST[value][1]
        # print(reduce(lambda x, y: x + y, tl))
        return reduce(lambda x, y: x + '|' + y, tl)

    @staticmethod
    def __paragraphs_to_sentences(paragraphs: str) -> 'List[str]':
        """
        先将段去空格等空格符，然后使用分句，返回单句list
        :return:
        """
        result = list()
        for para in paragraphs:
            para = para.replace(' ', '').replace('\r', '') \
                .replace('\n', '').replace('\t', '').replace('_', '')
            args = re.split(NLPProcess.PUNCTUATION_STR, para)
            result.extend(args)
        return result

    @staticmethod
    def paragraphs_to_sentences(paragraphs: str) -> 'List[str]':
        """
        先将段去空格等空格符，然后使用分句，返回单句list
        :return:
        """
        result = list()
        for para in paragraphs:
            para = para.replace(' ', '').replace('\r', '') \
                .replace('\n', '').replace('\t', '').replace('_', '')
            args = re.split(NLPProcess.PUNCTUATION_STR, para)
            result.extend(args)
        return result



    @staticmethod
    def sentences_to_words(sentences: 'List[str]') -> 'List[List[str]]':
        """
        将单句集合中的每个句子分词，并去除停用词和标点，得到单句分词的list的list
        :return:
        """
        data_set = list()
        for sent in sentences:
            words = jieba.cut(sent)
            words = list(filter(lambda x: x != '' and not NLPProcess.STOP_WORD_DIC.get(x), words))
            if words:
                data_set.append(words)
        return data_set

    @staticmethod
    def __train_word2vec_model(words: 'List[List[str]]'):
        """
        训练词向量的模型
        :param words:分词后的句子集合
        :return:
        """
        frequency = defaultdict(int)
        for text in words:
            for token in text:
                frequency[token] += 1
        precessed_corpus = [[token for token in text if frequency[token] > 1] for text in words]
        # print(precessed_corpus)
        model = word2vec.Word2Vec(precessed_corpus, min_count=0)
        # print(model['转让'])
        model.save(WORD2VEC_MODEL)
        return model

    @staticmethod
    def __words_to_vector(words: 'List[str]'):
        """
        将词库中的词计算成词向量
        :param words:
        :return:
        """
        if not os.path.exists(WORD2VEC_MODEL):
            return NLPProcess.__train_word2vec_model(words)
        else:
            return Word2Vec.load(WORD2VEC_MODEL)

    @staticmethod
    def slidings_to_words(slidings: 'List[List[str]]') -> 'List[List[str]]':
        """
        将单句集合中的每个句子分词，并去除停用词和标点，得到单句分词的list的list
        :return:
        """
        tmp = list()
        for sli in slidings:
            data_set = list()
            for sent in sli:
                words = jieba.cut(sent)
                words = list(filter(lambda x: x != '' and not NLPProcess.STOP_WORD_DIC.get(x), words))
                if words:
                    data_set.extend(words)
            tmp.append(data_set)
        return tmp

    @staticmethod
    def dump(path: 'str', obj: 'NLPProcess'):
        """
        将NlPProcess写入文件
        :param path:
        :param obj:
        :return:
        """
        with open(path, 'wb') as file:
            pickle.dump(obj, file)

    @staticmethod
    def load(path: str):
        """
        从文件中读取NlPProcess
        :param path:
        :return:
        """
        with open(path, 'rb') as file:
            return pickle.load(file)
