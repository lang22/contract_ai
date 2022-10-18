# -*- coding: utf-8 -*-
"""
Created on Mon Mar 25 18:52:32 2019

对于智能法审项目，将二进制存储成文件，再读取正文和表格，拼接成长字符串。

@author: lushilun
"""

import base64
import random
import docx
import subprocess
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
import pandas as pd
import numpy as np
import os


class readfile():
    def __init__(self):
        if not os.path.exists('temp'):
            os.makedirs('temp')
        pass

    def readmain(self, filedata):
        # 这是一个业务逻辑函数，对于每个项目，该函数可更改
        assert type(filedata) == dict

        if filedata['fileType'] == 'doc':
            filedata['docFile'] = self.readdoc(filedata['docFile'])
        elif filedata['fileType'] == 'docx':
            filedata['docFile'] = self.readdocx(filedata['docFile'])
        elif filedata['fileType'] == 'pdf':
            filedata['docFile'] = self.readpdf(filedata['docFile'])
        elif filedata['fileType'] in ['xls', 'xlsx']:
            filedata['docFile'] = self.readxls(filedata['docFile'])
        else:
            print('The type of file is not suitable! Please check it!')
            filedata['docFile'] = ''
        # 返回结果状态编码
        result = '200'
        if not filedata['docFile']:
            result = 'E001'
        return filedata, result

    def readdoc(self, content):
        # 将word文档存储，并读取内容
        try:
            filename = 'temp/' + str(random.randint(1, 10)) + '.doc'
            content = base64.b64decode((content))
            with open(filename, 'wb') as f:
                f.write(content)
            output = subprocess.check_output(["antiword", filename])
            output = bytes.decode(output)
            f.close()
        except Exception as e:
            print(e)
            output = ''
        finally:
            return output

    def readdocx(self, content):
        # 将docx文档存储，并读取内容
        try:
            filename = 'temp/' + str(random.randint(1, 10)) + '.docx'
            content = base64.b64decode((content))
            with open(filename, 'wb') as f:
                f.write(content)
            filecontent = docx.Document(filename)
            main_text = ''
            for paragraph in filecontent.paragraphs:
                # 对于正文来说，读取其内容
                main_text = main_text + '\n' + paragraph.text

            table_text = ''
            for table in filecontent.tables:
                # 对于每一个表格来说，读取其内容
                for i in range(len(table.rows)):
                    for j in range(len(table.columns)):
                        table_text = table_text + '\t' + table.cell(i, j).text
                table_text = table_text + '\n'
            output = main_text + table_text
            output = output.strip()
        except Exception as e:
            print(e)
            output = ''
        finally:
            print(output)
            return output

    def readpdf(self, content):
        # TODO(lushilun):将PDF格式的字符进行解码，并保存
        content = base64.b64decode((content))
        filename = 'temp/' + str(random.randint(1, 10)) + '.pdf'
        with open(filename, 'wb') as f:
            f.write(content)

        fp = open(filename, 'rb')
        # 用文件对象创建一个PDF文档分析器
        parser = PDFParser(fp)
        # 创建一个PDF文档
        doc = PDFDocument()
        # 连接分析器，与文档对象
        parser.set_document(doc)
        doc.set_parser(parser)

        # 提供初始化密码，如果没有密码，就创建一个空的字符串
        doc.initialize()

        # 检测文档是否提供txt转换，不提供就忽略
        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed
        else:
            # 创建PDF，资源管理器，来共享资源
            rsrcmgr = PDFResourceManager()
            # 创建一个PDF设备对象
            laparams = LAParams()
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            # 创建一个PDF解释其对象
            interpreter = PDFPageInterpreter(rsrcmgr, device)

            # 循环遍历列表，每次处理一个page内容
            # doc.get_pages() 获取page列表
            output = ''
            for page in doc.get_pages():
                interpreter.process_page(page)
                # 接受该页面的LTPage对象
                layout = device.get_result()
                # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
                # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
                # 想要获取文本就获得对象的text属性，
                for x in layout:
                    if (isinstance(x, LTTextBoxHorizontal)):
                        results = x.get_text()
                        output += results.strip('\n')
        return output

    def readxls(self, content):
        # TODO(lushilun): 将xls文件解码，并读取
        try:
            filename = 'temp/' + str(random.randint(1, 10)) + '.xls'
            content = base64.b64decode((content))
            with open(filename, 'wb') as f:
                f.write(content)
                filedata = pd.read_excel(filename, header=None)
            filedata = filedata.replace(np.nan, '')
            row_num, col_num = filedata.shape
            output = ''
            for i in range(row_num):
                for j in range(col_num):
                    output = output + '\t' + str(filedata.iloc[i, j])
                output = output + '\n'
            output = output.strip()
        except Exception as e:
            print(e)
            output = ''
        finally:
            return output

# import json
# files = readfile()
# with open("some.json") as f:
#     param = json.load(f)
# print(files.readmain(param))

if __name__ == '__main__':
    out_pp = readfile.readdoc('抵押协议采矿权-测试文档1.doc')
    print(out_pp)
