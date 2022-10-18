#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Time       : 2018/12/16 1:53
@Author     : jzs
@File       : public_substring.py
@Software   : PyCharm
@Description: ......
"""
from docx import Document

from app.tools.suffix_tree import create_suffix_tree, INF, print_suffix_tree

# 根节点的数据
ROOT_DATA = (0, 0)

# 目标叶子的起始坐标
TAG_DATA = (0, 0)

# 分割符号
SPLIT_FLAG = '#'

# 结束符
END_FLAG = '$'

# 公共子串的字典
__substring_dict = dict()


def __set_leaf_tag_data(tag1: int, tag2: int):
    """
    设置叶子查找目标

    :param tag1: 目标的下标开始
    :param tag2: 目标的下标结束
    :return:
    """
    global TAG_DATA
    TAG_DATA = (tag1, tag2)


def __is_public_substring(i: int, j: int):
    """
    判断是否是公共子串

    :param i: 开始坐标
    :param j: 结束坐标
    :return:
    """
    global TAG_DATA
    return i <= TAG_DATA[0] and j == TAG_DATA[1]


def __add_public_substring(children, data, ancestor):
    """
    将该叶子以及其祖先添加到字典中，用于过滤相同的祖先

    :param children: 结点的孩子
    :param data: 结点的数据
    :param ancestor: 结点的组先
    :return:
    """
    global __substring_dict
    if not children and len(ancestor) > 1 and __is_public_substring(*data):
        count = sum(j - i + 1 for i, j in ancestor[:-1])
        end = ancestor[-1][0] - 1
        start = end - count
        # print('start,end', (start,end))
        key = __substring_dict.get(end)
        if not key:
            __substring_dict[end] = [start, end]
        elif key[0] > start:
            __substring_dict[end][0] = start


def __post_order(root):
    """
    递归后根遍历字典树，并记录根结点的祖先，查找含有分割符号的叶子的组先

    :param root: 字典树的根
    :return:
    """
    data = (root[0], root[1])
    if data != ROOT_DATA:
        __post_order.ancestor.append(data)

    children = root[2].values()
    for child in children:
        __post_order(child)

    __add_public_substring(children, data, __post_order.ancestor)

    if __post_order.ancestor:
        __post_order.ancestor.pop()


def __pruning(root):
    """
    字典树的剪枝，减去第二层的叶子节点

    :param root: 字典树的根
    :return:
    """

    return root


def __public_substring_slices(string1: str, string2: str) -> 'list[tuple[int,int]]':
    """
    查找string1和string2的最长公共字串，返回基于string1的最长公共子串的切片list
    将string1和string2变成string1#string2&压入广义后缀树，后序遍历查找公共子串

    :param string1: 串1
    :param string2: 串1
    :return:
    """
    tmp = string1, SPLIT_FLAG, string2, END_FLAG
    string = str().join(tmp)
    root = create_suffix_tree(string)
    # print_suffix_tree(string)
    global __substring_dict
    __substring_dict = dict()
    __set_leaf_tag_data(string.find('#') + 1, INF)
    __post_order.ancestor = list()
    __post_order(root)  # 开始递归

    return __substring_dict.values()


def __get_LIS(arr: 'list[int]'):
    """
    最长递增子序列

    :param arr: 递增的序列
    :return: 得到最长递增子序列
    """
    n = len(arr)
    m = [0] * n
    for x in range(n - 2, -1, -1):
        for y in range(n - 1, x, -1):
            if arr[x] < arr[y] and m[x] <= m[y]:
                m[x] += 1
        max_value = max(m)
        result = []
        for i in range(n):
            if m[i] == max_value:
                result.append(arr[i])
                max_value -= 1
    return result


def get_public_substring_list(string1: str, string2: str) -> 'list[str]':
    """
    查找string1和string2的最长公共字串集合

    :param string1: 串1
    :param string2: 串2
    :return: 基于string1的最长公共字串list
    """

    slices = __public_substring_slices(string1, string2)
    slices = sorted(slices, key=lambda x: x[0])
    substrings = [string1[slice(*s)] for s in slices]

    return substrings


def is_increase_substrings(string1: str, string2: str, substrings: 'list[str]') -> bool:
    """
    判断这个基于string1的string1和string2的最长公共字串集合在string2中是否是递增序列

    :param string1:  串1
    :param string2:  串2
    :param substrings: 基于string1的最长公共字串list
    :return: 是否是递增序列
    """
    index1 = [string1.find(s) for s in substrings]
    index2 = [string2.find(s) for s in substrings]


if __name__ == '__main__':
    # sss = [('债权转让协议中，规定债权人必须要有义务执行该项任务。',
    #         '在债权转让协议中，要求债权人有责任执行该任务。'),
    #
    #        ('关于印发《不良资产包收购处置业务审查要点指引（试行）》',
    #         '关于《不良资产包的收购重组业务审查要点指引》'),
    #
    #        ('为规范公司不良资产包收购处置业务的审查工作',
    #         '为了规范公司业务的审查工作'),
    #
    #        ('不良资产包收购处置业务审查应采取全面审查与重点审查、现场核查与非现场核查相结合的方式',
    #         '不良资产包的收购处置业务的审查工作，应采取全面审查与重点审查、现场审查与非现场核查相结合的方法')]
    #
    # for s1, s2 in sss:
    #     print('s1:', s1)
    #     print('s2:', s2)
    #     substrings = get_public_substring_list(s1, s2)
    #     print(substrings)

    document = Document('D:/合同机器人项目相关/测试预测用文档kkk.docx')
    print([p.text for p in document.paragraphs])
    for p in document.paragraphs:
        print(p.text)
