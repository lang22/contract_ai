#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Time       : 2018/12/13 13:51
@Author     : jzs
@File       : docx_tools.py
@Software   : PyCharm
@Description: docx的一些工具
"""
import copy

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def __add_text(paragraph, font, text):
    """
    在段落中添加文字

    :param paragraph: 段落
    :param text: 文字
    :return:
    """
    run = paragraph.add_run(text)
    run.font.name = font.name
    run.font.size = font.size
    run.font.bold = font.bold
    run.font.highlight_color = font.highlight_color
    return run


def __add_highlight_color_text(paragraph, font, text, highlight_color):
    """
    在段落中添加高亮文字

    :param paragraph: 段落
    :param text: 文字
    :return:
    """
    run = paragraph.add_run(text)
    run.font.name = font.name
    run.font.size = font.size
    run.font.bold = font.bold
    run.font.highlight_color = highlight_color
    return run


def __confirm_runs(paragraph):
    """
    将runs赋值给每个字，只赋值一次

    :param paragraph: 段
    :return:
    """
    if len(paragraph.text) + 1 != len(paragraph.runs):
        font = paragraph.runs[0].font
        text = paragraph.text
        paragraph.text = ''
        for c in text:
            __add_text(paragraph, font, c)


def set_paragraph_text_highlight_by_slice(paragraph,
                                          para_slice,
                                          highlight_color: str = 'YELLOW') -> None:
    """
    使用字符串切片，设置一段高亮或一段中某些文字的高亮

    :param paragraph: 段落
    :param para_slice 需要标高亮的切片
    :param highlight_color: WD_COLOR_INDEX中的颜色名字字符串，默认为YELLOW
    :return:
    """
    if not highlight_color:
        return

    s1, s2 = para_slice
    print("s2 - s1 == len(paragraph.text) : ", s2 - s1, len(paragraph.text), paragraph.text)
    if s2 - s1 == len(paragraph.text):
        for runs in paragraph.runs:
            runs.font.highlight_color = getattr(WD_COLOR_INDEX, highlight_color)
        return

    __confirm_runs(paragraph)

    for i in range(para_slice[0] + 1, para_slice[1] + 1):
        print('WD_COLOR_INDEX', WD_COLOR_INDEX, 'highlight_color', highlight_color)
        if i >= len(paragraph.runs):
            continue
        paragraph.runs[i].font.highlight_color = getattr(WD_COLOR_INDEX, highlight_color)


def __find_tag_para_insert_index(tag_para,
                                 source_para,
                                 source_para_slice):
    """
    找到tag_para串中对应的source_para_slice字符串的插入下标

    :param tag_para: 待审核文档段落
    :param source_para: 模板文档的段落
    :param source_para_slice: 模板文档切片
    :return:
    """
    ss1, ss2 = source_para_slice
    if ss1 == 0:
        return 0
    elif ss2 == len(source_para.text):
        return len(tag_para.text)

    print("del_string:", source_para.text[ss1: ss2])
    print('find_test_string:', source_para.text[ss1 - 1] + source_para.text[ss2])
    find_index = tag_para.text.find(source_para.text[ss1 - 1] + source_para.text[ss2])
    return find_index + 1


def insert_paragraph_text_highlight_by_slice(tag_para,
                                             source_para,
                                             source_para_slice,
                                             highlight_color: str = 'PINK') -> None:
    """
    使用字符串切片，插入一段高亮或一段中某些文字的高亮

    :param tag_para: 待审核文档段落
    :param source_para: 模板文档的段落
    :param source_para_slice: 模板文档切片
    :param highlight_color: WD_COLOR_INDEX中的颜色名字字符串，默认为PINK
    :return:
    """

    split_index = __find_tag_para_insert_index(tag_para, source_para, source_para_slice)
    ss1, ss2 = source_para_slice

    print("tag_para.text:", tag_para.text)
    __confirm_runs(tag_para)

    text_list = list(tag_para.text)
    runs_list = list(tag_para.runs)

    # 按添加位置的分割的段落前半部分字符list和run list
    text_list_front = text_list[0:split_index]
    runs_list_front = runs_list[1:split_index + 1]

    # 按添加位置的分割的后半部分字符list和run list
    text_list_tail = text_list[split_index:]
    runs_list_tail = runs_list[split_index + 1:]

    # 新添加的字符
    new_text = list(source_para.text[ss1:ss2])
    new_runs = source_para.runs[0]

    tag_para.text = ''
    # 处理前半部分
    for i, t in enumerate(text_list_front):
        __add_text(tag_para, runs_list_front[i].font, t)

    # 处理新添加
    for i, t in enumerate(new_text):
        __add_highlight_color_text(tag_para, new_runs.font, t, getattr(WD_COLOR_INDEX, highlight_color))

    # 处理后部分
    for i, t in enumerate(text_list_tail):
        __add_text(tag_para, runs_list_tail[i].font, t)


def add_paragraph_text_highlight_before(before_paragraph,
                                        model_para,
                                        text: str,
                                        highlight_color: str = 'PINK') -> None:
    """
    在段落before_paragraph前添加一段与段落before_paragraph样式一致的高亮的段落


    :param before_paragraph: 需要插入的段落之前的段落
    :param model_para:  模板文档段落
    :param text: 添加的文字
    :param highlight_color: 高亮颜色
    :return:
    """
    new_para = before_paragraph.insert_paragraph_before("")
    new_runs = model_para.runs[0]
    __add_highlight_color_text(new_para, new_runs.font, text, getattr(WD_COLOR_INDEX, highlight_color))


if __name__ == '__main__':
    d1 = Document('C:/Users/11378/Desktop/TMP/test.docx')
    d2 = Document('C:/Users/11378/Desktop/TMP/model.docx')

    set_paragraph_text_highlight_by_slice(d1.paragraphs[46], (28, 30))

    del1 = (4, 8, 8)
    del2 = (7, 8, 46)
    del3 = (27, 28, 46)

    insert_paragraph_text_highlight_by_slice(d1.paragraphs[8], d2.paragraphs[8], (4, 8))
    set_paragraph_text_highlight_by_slice(d1.paragraphs[46], (35, 38))

    insert_paragraph_text_highlight_by_slice(d1.paragraphs[46], d2.paragraphs[46], (7, 8))
    insert_paragraph_text_highlight_by_slice(d1.paragraphs[46], d2.paragraphs[46], (27, 28))

    d1.save('C:/Users/11378/Desktop/测试-可以.docx')
