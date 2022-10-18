"""
@author: Chenwei YAN
@Created: 2018/5/11
"""
from docx import Document
from docx.shared import *
import re
import os
import pandas as pd


def compare(upload_path_userup):
    # filename = '形式化审核.docx'
    # basepath = os.path.dirname(__file__)
    # upload_path = os.path.join(basepath, 'static/docx_repo/docx_rd/', filename)
    # print(upload_path)
    template_list = read_template(upload_path_userup)

    actual_list, docx_file = read_file(upload_path_userup)
    status, new_num = check(template_list, actual_list)
    # template_list =
    df_table = to_csv(template_list, status, new_num)
    # return df_table, docx_file
    return df_table


def read_template(filename1):
    template = Document(filename1)
    str_list = []
    paranum_list = []
    para_num_list = []
    # print('\n' + filename1 + ' 中:')
    para_num = 0
    for para in template.paragraphs:
        if re.match(r' *★*第(.*)条', para.text):

            paranum_list.append(para_num)
            n = para.text.split(' ')
            # print(para.text)
            # print(template.paragraphs[para_num-2].text)
            para_num_list.append(para_num)
            for i in range(len(n)):
                if n[i] != '':
                    str_list.append(n[i])
        para_num += 1
    # print(str_list)
    return str_list


def read_file(filename1):
    template = Document(filename1)
    str_list = []
    paranum_list = []
    para_num_list = []
    # print('\n' + filename1 + ' 中:')
    para_num = 0
    for para in template.paragraphs:
        if re.match(r' *★*第(.*)条', para.text):

            paranum_list.append(para_num)
            paragraph = template.paragraphs[para_num]
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(205, 0, 0)
                run.font.bold = True
            # filename = 'xsh_exam.docx'
            # basepath = os.path.dirname(__file__)
            # upload_path = os.path.join(basepath, 'static/docx_repo/docx_gt/', filename)
            # template.save(upload_path)
            # template.save('D:/jupyter_datafile/docx_temp/xsh_exam.docx')

            n = para.text.split(' ')
            # print(para.text)
            # print(template.paragraphs[para_num-2].text)
            para_num_list.append(para_num)
            for i in range(len(n)):
                if n[i] != '':
                    str_list.append(n[i])
        para_num += 1
    # print(str_list)
    # cont = template.paragraphs
    # for item in cont:
    #     print(item.text)
    return str_list, template


def check(list1, list2):
    check_list = ['在合同中未找到该条'] * int(len(list1) / 2)
    new_num_list = [' '] * int(len(list1) / 2)
    # print('\n')
    if len(list1) > len(list2):
        len_list = len(list2)
    else:
        len_list = len(list1)

    for i in range(1, len_list, 2):
        if list1[i - 1] == list2[i - 1] and list1[i] == list2[i]:
            check_list[int(i / 2)] = '一致'
            new_num_list[int(i / 2)] = '通过'
        # elif list1[i] in list2:
        if list1[i - 1] == list2[i - 1] and list1[i] != list2[i]:
            for j in range(1, len_list, 2):
                if re.search(list1[i], list2[j]) != None:
                    # print(list1[i-1] + '已填写')
                    check_list[int(i / 2)] = '不一致'
                    # for v in range(1,len_list,2):
                    #     if re.search(list1[v],list2[v]) != None:
                    # if list2[j] == list1[i]:
                    new_num_list[int(i / 2)] = '对应合同中' + list2[j - 1]
            #
        # if re.search(list1[i],list2[i]) == None:
        #     # print('该条目在合同中不存在：' + list1[i-1] + ' ' + list1[i])
        #     check_list[int(i / 2)] = '在合同中未找到该条'

    return check_list, new_num_list


def to_csv(list1, check_status, new_num_list):
    # 生成DF
    num1 = []
    value1 = []
    for i in range(0, len(list1), 2):
        num1.append(list1[i])
        value1.append(list1[i + 1])
    c = {"T_item": num1, "T_text": value1, 'Check_stata': check_status, 'Othe': new_num_list}
    df = pd.DataFrame(c)
    df.index = range(1, len(df) + 1)
    # print('\n')
    # print(df)
    return df

# if __name__ == '__main__':
#     compare('/Users/yanchenwei/Desktop/1 中融-宏金18号项目单一资金信托之信托贷款合同-补充协议 法律V1.docx')
