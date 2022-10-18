#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
@Time       : 2018/11/22 12:42
@Author     : jzs
@File       : flow_table_dump.py
@Software   : PyCharm
@Description: ......
"""
import os
import time

from docx import Document
from docx.shared import Pt
from flask import request

from app.tools import random_fliename
from config import Config

flow_table1 = os.path.join(Config.CONT_FLOW_TABLE_DIR_PATH, '合同面签记录表.docx')
flow_table2 = os.path.join(Config.CONT_FLOW_TABLE_DIR_PATH, '合同审查申请表.docx')


def dump_flow_table(pt_name='', doc_date='', doc_loc='',
                    rows=()):
    """
    保存合同流转单1：合同面签记录表, 返回保存的文件的所在文件夹的绝对路径和该文件名
    :param pt_name:
    :param doc_date:
    :param doc_loc:
    :param rows:
    :return:
    """
    document = Document(flow_table1)
    paragraph = document.paragraphs[3]
    paragraph.text = ""
    run = paragraph.add_run(u'本记录表在' + doc_date + '于' + doc_loc + '填写。 ')

    run.font.size = Pt(14)
    table = document.tables[0]
    table.rows[0].cells[1].text = pt_name

    for i, name, cont, othe in rows:
        table.rows[i + 1].cells[1].text = name if name else ''
        table.rows[i + 1].cells[2].text = cont if cont else ''
        table.rows[i + 1].cells[3].text = othe if othe else ''

    filename = '合同面签记录表' + random_fliename() + '.docx'
    upload_path_temp = os.path.join(Config.CONT_FLOW_TABLE_DIR_PATH, filename)
    document.save(upload_path_temp)

    return Config.CONT_FLOW_TABLE_DIR_PATH, filename


def dump_flow_table2(dp_name='', pt_name='', pt_cont='', its=[[]]):
    """
    填写合同流转单2：合同审查申请表, 返回保存的文件的所在文件夹的绝对路径和该文件名
    :return:
    """
    document = Document(flow_table2)
    table = document.tables[0]
    table.rows[0].cells[1].text = dp_name
    table.rows[1].cells[1].text = pt_name
    table.rows[2].cells[1].text = pt_cont

    for i in range(4):
        it = its[i]
        for j in range(len(it)):
            table.rows[i + 4].cells[j].text = it[j] if it[j] else ''

    filename = '合同审查申请表' + random_fliename() + '.docx'
    upload_path_temp = os.path.join(Config.CONT_FLOW_TABLE_DIR_PATH, filename)
    document.save(upload_path_temp)

    return Config.CONT_FLOW_TABLE_DIR_PATH, filename
